from __future__ import annotations

from pathlib import Path

from app.parsers.models import ParsedDocument


def parse_word(
    path: str | Path,
    *,
    url: str,
    source_id: str = "",
    content_hash: str = "",
    fetched_at: str = "",
    raw_path: str = "",
) -> ParsedDocument:
    path = Path(path)
    if path.suffix.lower() == ".doc":
        return ParsedDocument(
            source_id=source_id,
            url=url,
            title=path.name,
            text="Legacy .doc is not parsed by default. Optional conversion via LibreOffice can be added in deployment.",
            content_hash=content_hash,
            fetched_at=fetched_at,
            raw_path=str(path),
            parser="word",
            metadata={"unsupported": ".doc"},
        )
    try:
        from docx import Document

        document = Document(str(path))
        parts = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]
        for table in document.tables:
            for row in table.rows:
                parts.append(" | ".join(cell.text.strip() for cell in row.cells))
        text = "\n".join(parts)
    except Exception as exc:
        text = f"[docx extraction failed: {exc}]"
    return ParsedDocument(
        source_id=source_id,
        url=url,
        title=path.name,
        text=text,
        content_hash=content_hash,
        fetched_at=fetched_at,
        raw_path=str(path),
        parser="word",
    )

